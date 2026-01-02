"""
K-Stay Document Service
Word 문서 생성 및 매핑 - DocumentProcessor 엔진 적용
섹션 기반 구조 지원 (self vs other_xxx)
TABLE_ROWS 전략 추가 (동적 테이블 행 입력)
"""

import streamlit as st
from typing import Dict, List, Optional, Any
from datetime import datetime
import io
import zipfile
import os
import re
import tempfile

# python-docx
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
    print("✅ python-docx 로드 성공")
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"⚠️ python-docx 미설치 - pip install python-docx 실행 필요")
except Exception as e:
    DOCX_AVAILABLE = False
    print(f"⚠️ python-docx 로드 실패: {e}")


# =============================================================================
# 성별 체크박스 매핑 (공통)
# =============================================================================

NATURALIZATION_CHECKBOX_MAP = {
    # 일반귀화
    "general_permanent_resident": ["「민법」상 성년이며 영주자격(F5)을 가지고 있는 사람"],
    
    # 간이귀화
    "simplified_parent_korean": ["부 또는 모가 대한민국의 국민이었던 사람"],
    "simplified_born_in_korea": ["대한민국에서 출생한 사람으로서 부 또는 모가 대한민국에서 출생한 사람"],
    "simplified_adopted": ["대한민국 국민의 양자(養子)로서 입양 당시 대한민국의 「민법」상 성년이었던 사람"],
    
    # 혼인귀화
    "marriage_2years": ["배우자와 혼인한 상태로 대한민국에 2년 이상 거주한 사람"],
    "marriage_3years_1year": ["배우자와 혼인한 후 3년이 지나고 혼인한 상태로 대한민국에 1년 이상 거주한 사람"],
    "marriage_spouse_unavailable": ["배우자의 사망", "실종 그 밖에 자신에게 책임이 없는 사유로 혼인생활 유지가 불가한"],
    "marriage_raising_child": ["배우자와의 혼인에 따라 출생한 미성년의 자녀를 양육하고 있거나 양육할 사람"],
    
    # 특별귀화
    "special_minor_adoptee": ["부 또는 모가 대한민국의 국민인 사람, 입양 당시 「민법」상 미성년이었던 사람"],
    "special_merit": ["대한민국에 특별한 공로가 있는 사람"],
    "special_merit_independence": ["독립유공자"],
    "special_merit_national": ["국가유공자"],
    "special_merit_national_interest": ["국익기여자"],
    "special_excellence": ["과학", "경제", "문화", "체육 등 특정 분야에서 매우 우수한 능력을 보유한 사람"],
}

GENDER_CHECKBOX_MAP = {
    "Male": ["남", "Male", "M", "Man"],
    "M": ["남", "Male", "M", "Man"],
    "남": ["남", "Male", "M", "Man"],
    "Female": ["여", "Female", "F", "Woman"],
    "F": ["여", "Female", "F", "Woman"],
    "여": ["여", "Female", "F", "Woman"],
}

YES_NO_CHECKBOX_MAP = {
    "Yes": ["예", "Yes", "Y", "있음"],
    "예": ["예", "Yes", "Y", "있음"],
    "No": ["아니오", "No", "N", "없음"],
    "아니오": ["아니오", "No", "N", "없음"],
}

VALUE_MAPS = {
    "GENDER": GENDER_CHECKBOX_MAP,
    "YES_NO": YES_NO_CHECKBOX_MAP,
}


# =============================================================================
# DocumentProcessor 클래스
# =============================================================================

class DocumentProcessor:
    """
    문서 처리 엔진 (중첩 테이블 지원 강화)

    """
       
    def _execute_hierarchical_checkbox(self, table, r_idx: int, c_idx: int,
                                        value: Any, field: Dict, 
                                        current_cell, row_cells):
        """
        계층적 체크박스 처리 (귀화 유형)
        
        귀화 유형은 문서에서 다음과 같은 구조를 가짐:
        | 일반귀화     | [  ] 「민법」상 성년이며... |
        | 간이귀화     | [  ] 부 또는 모가...        |
        |              | [  ] 대한민국에서 출생한... |
        """
        checkbox_config = field.get('checkbox_config', {})
        target_value = checkbox_config.get('value')
        is_nested = checkbox_config.get('is_nested', False)
        
        # 사용자가 선택한 값과 이 필드의 값이 일치하는지 확인
        if value != target_value:
            return  # 이 체크박스가 아니면 스킵
        
        # 체크박스 텍스트 찾기
        anchor_texts = field.get('anchor_text', [])
        if isinstance(anchor_texts, str):
            anchor_texts = [anchor_texts]
        
        self._log(f"      🔍 [계층적 체크박스] 검색 시작: {target_value}")
        
        # 현재 셀과 인접 셀들을 스캔
        checked = self._find_and_check_hierarchical_checkbox(
            table, r_idx, c_idx, anchor_texts, current_cell, row_cells
        )
        
        if checked:
            self._log(f"      ✅ [체크 완료] {target_value}")
        else:
            self._log(f"      ⚠️ [체크 실패] 체크박스를 찾을 수 없음: {anchor_texts}")
    
    def _find_and_check_hierarchical_checkbox(self, table, r_idx: int, c_idx: int,
                                               anchor_texts: List[str],
                                               current_cell, row_cells) -> bool:
        """
        계층적 체크박스를 찾아서 체크
        
        Returns:
            bool: 체크 성공 여부
        """
        # 스캔 대상 셀들 수집
        scan_targets = []
        
        # 1. 현재 행의 모든 셀
        for i, cell in enumerate(row_cells):
            scan_targets.append((f"현재행-C{i}", cell))
        
        # 2. 위/아래 행의 셀들
        for row_offset in [-2, -1, 1, 2]:
            target_row = r_idx + row_offset
            if 0 <= target_row < len(table.rows):
                try:
                    adjacent_cells = table.rows[target_row].cells
                    for k, cell in enumerate(adjacent_cells):
                        pos_name = f"행{'+' if row_offset > 0 else ''}{row_offset}-C{k}"
                        scan_targets.append((pos_name, cell))
                except:
                    continue
        
        # 각 셀에서 체크박스 찾기
        for pos_name, cell in scan_targets:
            original_text = cell.text
            if not original_text.strip():
                continue
            
            # 앵커 텍스트 매칭 확인
            all_anchors_found = all(
                anchor.lower() in original_text.lower() or
                self.normalize_text(anchor) in self.normalize_text(original_text)
                for anchor in anchor_texts
            )
            
            if not all_anchors_found:
                continue
            
            self._log(f"        📍 [발견] {pos_name}: '{original_text[:50]}...'")
            
            # 체크박스 패턴 찾아서 체크
            # 다양한 체크박스 패턴 지원
            patterns = [
                (r'\[\s*\]', '[V]'),           # [ ] -> [V]
                (r'□', '☑'),                   # □ -> ☑
                (r'☐', '☑'),                   # ☐ -> ☑
                (r'\(\s*\)', '(V)'),           # ( ) -> (V)
            ]
            
            new_text = original_text
            checked = False
            
            for pattern, replacement in patterns:
                if re.search(pattern, new_text):
                    # 해당 앵커 텍스트 근처의 체크박스만 체크
                    for anchor in anchor_texts:
                        # 앵커 텍스트 앞의 체크박스 찾기
                        anchor_pattern = fr'({pattern})(\s*)({re.escape(anchor[:20])})'
                        match = re.search(anchor_pattern, new_text, re.IGNORECASE)
                        if match:
                            new_text = re.sub(anchor_pattern, 
                                            fr'{replacement}\2\3', 
                                            new_text, count=1, flags=re.IGNORECASE)
                            checked = True
                            break
                    
                    # 앵커 텍스트 앞에서 못 찾으면, 줄 시작 부분의 체크박스 찾기
                    if not checked:
                        new_text = re.sub(pattern, replacement, new_text, count=1)
                        checked = True
                    
                    break
            
            if checked:
                cell.text = new_text
                self._log(f"        ✅ [체크됨] {pos_name}")
                return True
        
        return False
    
    def _execute_checkbox_with_value(self, table, r_idx: int, c_idx: int,
                                      value: Any, field: Dict,
                                      current_cell, row_cells, data: Dict):
        """
        값 입력이 필요한 체크박스 처리 (수반취득)
        
        예: "만 19세 미만의 자녀 (   )명에 대하여..."
        -> 체크박스 체크 + 괄호 안에 숫자 입력
        """
        checkbox_config = field.get('checkbox_config', {})
        value_field = checkbox_config.get('value_field')
        value_placeholder = checkbox_config.get('value_placeholder', '(   )')
        
        # 체크박스 활성화 여부 확인
        if not value or value == False or value == "No":
            return
        
        # 값 필드에서 실제 값 가져오기
        actual_value = data.get(value_field, '')
        
        anchor_texts = field.get('anchor_text', [])
        if isinstance(anchor_texts, str):
            anchor_texts = [anchor_texts]
        
        self._log(f"      🔍 [값 체크박스] 검색: {anchor_texts}, 값: {actual_value}")
        
        # 스캔 대상 셀들
        scan_targets = [(f"현재행-C{i}", cell) for i, cell in enumerate(row_cells)]
        
        for row_offset in [-1, 1]:
            target_row = r_idx + row_offset
            if 0 <= target_row < len(table.rows):
                try:
                    adjacent_cells = table.rows[target_row].cells
                    for k, cell in enumerate(adjacent_cells):
                        scan_targets.append((f"행{'+' if row_offset > 0 else ''}{row_offset}-C{k}", cell))
                except:
                    continue
        
        for pos_name, cell in scan_targets:
            original_text = cell.text
            if not original_text.strip():
                continue
            
            # 앵커 텍스트 매칭
            if not any(anchor in original_text for anchor in anchor_texts):
                continue
            
            self._log(f"        📍 [발견] {pos_name}")
            
            new_text = original_text
            
            # 1. 체크박스 체크
            checkbox_patterns = [
                (r'\[\s*\]', '[V]'),
                (r'□', '☑'),
                (r'☐', '☑'),
            ]
            
            for pattern, replacement in checkbox_patterns:
                if re.search(pattern, new_text):
                    new_text = re.sub(pattern, replacement, new_text, count=1)
                    break
            
            # 2. 괄호 안에 값 입력
            if actual_value:
                # (   ) 또는 (  ) 패턴을 찾아서 값으로 대체
                placeholder_pattern = r'\(\s*\)'
                if re.search(placeholder_pattern, new_text):
                    new_text = re.sub(placeholder_pattern, f'( {actual_value} )', new_text, count=1)
            
            if new_text != original_text:
                cell.text = new_text
                self._log(f"        ✅ [완료] 체크 + 값 입력: {actual_value}")
                return
        
        self._log(f"      ⚠️ [실패] 체크박스를 찾을 수 없음")

    
    def __init__(self, user_data: Dict[str, Any], form_data: Dict[str, Any], narrative_data: Dict[str, Any]):
        self.user_data = user_data or {}
        self.form_data = form_data or {}
        self.narrative_data = narrative_data or {}
        self.logs = []
        self._prepare_derived_data()
    def _apply_field(self, all_tables, field: Dict, data: Dict, target: str):
        """
        일반 필드(단일 값) 매핑 처리
        문서 전체 테이블을 순회하며 라벨(anchor_text)을 찾고 전략(strategy)을 실행
        """
        # [수정 1] 호환성 확보: anchor_text(구버전)와 key_text(신버전) 둘 다 확인
        anchors = field.get('anchor_text') or field.get('key_text')
        data_key = field.get('data_key')
        strategy = field.get('strategy', 'NEXT_CELL')
        value = data.get(data_key)
        if value is None or value == "":
            return 
        # HIERARCHICAL_CHECKBOX 특별 처리
        if strategy == "HIERARCHICAL_CHECKBOX":
            checkbox_config = field.get('checkbox_config', {})
            target_value = checkbox_config.get('value')
            
            # 사용자가 선택한 값과 이 필드의 타겟 값이 다르면 스킵
            if value != target_value:
                return
            
            self._log(f"    🎯 [계층적 체크박스] {data_key} = {value}")
            
            # 모든 테이블에서 해당 체크박스 찾기
            anchors = field.get('anchor_text', [])
            if isinstance(anchors, str):
                anchors = [anchors]
            
            for table in all_tables:
                for r_idx, row in enumerate(table.rows):
                    try:
                        row_cells = row.cells
                    except:
                        continue
                    
                    for c_idx, cell in enumerate(row_cells):
                        cell_text = cell.text
                        
                        # 앵커 텍스트가 포함된 셀 찾기
                        if any(anchor in cell_text for anchor in anchors):
                            self._execute_hierarchical_checkbox(
                                table, r_idx, c_idx, value, field, cell, row_cells
                            )
                            return  # 하나 찾으면 종료

        # 1. 데이터

        # 2. 검색할 라벨 텍스트 확인
        # [수정 2] 리스트 형태와 문자열 형태 모두 처리
        if isinstance(anchors, str):
            anchors = [anchors]
        
        if not anchors:
            return

        # 검색어 정규화
        normalized_anchors = [self.normalize_text(a) for a in anchors]
        
        # 3. 미리 수집된 모든 테이블(all_tables) 순회
        # (기존 코드처럼 함수 안에서 다시 _get_all_tables를 호출하면 매우 느려짐)
        for table in all_tables:
            for r_idx, row in enumerate(table.rows):
                try:
                    row_cells = row.cells
                except:
                    continue

                for c_idx, cell in enumerate(row_cells):
                    cell_text_clean = self.normalize_text(cell.text)

                    # 라벨 매칭 (List 중 하나라도 포함되면 OK)
                    if any(norm_a in cell_text_clean for norm_a in normalized_anchors):
                        self._log(f"    🎯 [매칭] '{anchors[0]}' -> 값: {str(value)[:10]}... (T?-R{r_idx}-C{c_idx})")
                        
                        # 전략 실행
                        self._execute_strategy(
                            strategy, table, r_idx, c_idx, 
                            value, field, cell, row_cells
                        )
                        return # 하나 찾으면 해당 필드 처리 종료
        
        self._log(f"    ⚠️ [실패] 문서에서 라벨을 찾을 수 없음: '{anchors[0]}'")
    def _prepare_derived_data(self):
        # (기존과 동일)
        if self.user_data.get('surname') and self.user_data.get('given_name'):
            if 'full_name' not in self.user_data:
                self.user_data['full_name'] = f"{self.user_data['surname']} {self.user_data['given_name']}"
        if 'application_date' not in self.form_data:
            self.form_data['application_date'] = datetime.now().strftime('%Y.%m.%d')

    def get_data_for_section(self, target: str) -> Dict[str, Any]:
        # (기존과 동일)
        if target == "self":
            merged = {**self.user_data, **self.narrative_data}
            for key, value in self.form_data.items():
                if not any(key.startswith(p) for p in ['guarantor_', 'spouse_', 'inviter_', 'employer_']):
                    merged[key] = value  # [수정 완료] 'a=' 를 '=' 로 변경
            return merged
        else:
            from config.settings import TARGET_INFO
            target_info = TARGET_INFO.get(target, {})
            prefix = target_info.get("prefix", "")
            if not prefix: return self.form_data
            return {k: v for k, v in self.form_data.items() if k.startswith(prefix)}

    def normalize_text(self, text: str) -> str:
        return re.sub(r'\s+', '', text).lower()
    
    def _log(self, message: str):
        self.logs.append(message)
        print(message)

    # --- [핵심] 모든 테이블(중첩 포함) 수집 메서드 ---
# --- [핵심] 모든 테이블(중첩 포함) 수집 메서드 ---
    def _get_all_tables(self, doc_or_cell):
        """문서 내의 모든 테이블을 재귀적으로 찾아서 리스트로 반환"""
        tables = []
        
        # 1. 현재 객체(Document 또는 Cell)에 속한 테이블 수집
        if hasattr(doc_or_cell, 'tables'):
            current_tables = doc_or_cell.tables
            for table in current_tables:
                tables.append(table)
                # 2. 각 테이블의 셀 내부를 재귀적으로 탐색
                for row in table.rows:
                    for cell in row.cells:
                        tables.extend(self._get_all_tables(cell))
        return tables

    def process_file(self, mapping_config: Dict, input_path: str, output_path: str) -> bool:
        if not DOCX_AVAILABLE: return False
        
        try:
            doc = Document(input_path)
            self._log(f"🔄 [분석 시작] {os.path.basename(input_path)}")
            
            # 1. 문서 내 모든 테이블 수집
            all_tables = self._get_all_tables(doc)
            self._log(f"📊 감지된 총 테이블 수(중첩 포함): {len(all_tables)}개")

            sections = mapping_config.get('sections', [])
            if sections:
                for section in sections:
                    section_name = section.get('section_name', '')
                    target = section.get('target', 'self')
                    fields = section.get('fields', [])
                    
                    # -------------------------------------------------------
                    # [Step 1] 검색 범위(Scope) 설정: (Table, Start_Row_Index)
                    # -------------------------------------------------------
                    target_scopes = []
                    table_match_text = section.get('table_match_text')
                    
                    if table_match_text:
                        # 텍스트 정규화 (공백/줄바꿈 제거하여 비교)
                        clean_match_text = self.normalize_text(table_match_text)
                        
                        found_count = 0
                        for table in all_tables:
                            found_in_table = False
                            # 테이블의 모든 행을 검사하여 시작점(Row Index) 찾기
                            for r_idx, row in enumerate(table.rows):
                                # 행 전체 텍스트 합치기 (셀 병합 고려)
                                row_text = "".join([c.text for c in row.cells])
                                if clean_match_text in self.normalize_text(row_text):
                                    # 해당 텍스트가 포함된 행부터 검색 시작
                                    target_scopes.append((table, r_idx))
                                    found_in_table = True
                                    found_count += 1
                                    self._log(f"   📍 [Scope] '{section_name}' 시작점 발견: Table(Row {r_idx}) - '{table_match_text}'")
                                    break # 한 테이블 내에서 시작점은 한 번만 찾음
                            
                            # (옵션) 테이블 텍스트에는 있지만 행 매칭 실패 시(구조적 문제 등)
                            # 필요하다면 여기서 전체 테이블 추가 로직 구현 가능
                        
                        if not target_scopes:
                            self._log(f"   ⚠️ [경고] '{section_name}' 섹션의 기준 텍스트 '{table_match_text}'를 찾지 못했습니다. 건너뜁니다.")
                            continue
                        else:
                            self._log(f"   🔒 [Scope] '{section_name}' 섹션은 {found_count}개의 지점으로 범위 제한됨")
                            
                    else:
                        # 매칭 텍스트가 없으면: 모든 테이블, 0번 줄부터 검색
                        target_scopes = [(t, 0) for t in all_tables]

                    # -------------------------------------------------------
                    # [Step 2] 데이터 준비 및 필드 매핑 실행
                    # -------------------------------------------------------
                    self._log(f"\n📁 [섹션 처리] {section_name}")
                    section_data = self.get_data_for_section(target)
                    all_data = {**section_data, **self.form_data}
                    
                    for field in fields:
                        strategy = field.get('strategy', 'NEXT_CELL')
                        
                        if strategy == "TABLE_ROWS":
                            # [Code 1 복원] TABLE_ROWS 로직 실행
                            # 주: TABLE_ROWS는 보통 전체 테이블 구조를 보고 동작하므로 all_tables를 넘기거나,
                            # 필요하다면 target_scopes에 있는 테이블들만 추려서 넘길 수도 있습니다.
                            # 여기서는 기존 호환성을 위해 all_tables를 사용합니다.
                            self._apply_table_rows(all_tables, field, all_data)
                            
                        else:
                            # [Code 2 적용] Scope(테이블+시작행) 정보를 이용한 필드 적용
                            # _apply_field 대신 _apply_field_scoped를 사용해야 정확한 행부터 찾습니다.
                            if hasattr(self, '_apply_field_scoped'):
                                self._apply_field_scoped(target_scopes, field, section_data)
                            else:
                                # 혹시 _apply_field_scoped 메서드가 아직 없다면 기존 방식 호환
                                # (하지만 정확도는 떨어질 수 있음)
                                filtered_tables = [t for t, r in target_scopes]
                                self._apply_field(filtered_tables, field, section_data, target)
            
            doc.save(output_path)
            self._log(f"✅ [저장 완료] {output_path}")
            return True
            
        except Exception as e:
            self._log(f"❌ 오류: {e}")
            import traceback
            traceback.print_exc()
            return False
    def _apply_field_scoped(self, target_scopes, field: Dict, data: Dict):
        """
        범위(Scope)가 지정된 필드 처리
        target_scopes: List[(Table, Start_Row_Index)]
        """
        anchors = field.get('anchor_text') or field.get('key_text')
        data_key = field.get('data_key')
        strategy = field.get('strategy', 'NEXT_CELL')
        value = data.get(data_key)
        
        if value is None or value == "": return
        if isinstance(anchors, str): anchors = [anchors]
        if not anchors: return

        normalized_anchors = [self.normalize_text(a) for a in anchors]

        # 지정된 범위(테이블 + 시작행)만 순회
        for table, start_row in target_scopes:
            # start_row 부터 끝까지만 반복
            for r_idx in range(start_row, len(table.rows)):
                row = table.rows[r_idx]
                try:
                    row_cells = row.cells
                except:
                    continue

                for c_idx, cell in enumerate(row_cells):
                    cell_text_clean = self.normalize_text(cell.text)
                    
                    # 앵커 찾기
                    if any(norm_a in cell_text_clean for norm_a in normalized_anchors):
                        self._log(f"    🎯 [매칭] '{anchors[0]}' (Row {r_idx}) -> 값 입력")
                        
                        # 전략 실행 (기존 메서드 활용)
                        self._execute_strategy(
                            strategy, table, r_idx, c_idx, 
                            value, field, cell, row_cells
                        )
                        return # 하나 찾으면 종료

    def _apply_table_rows(self, all_tables, field: Dict, data: Dict):
        """
        TABLE_ROWS 전략 최종본
        기능 1: must_contain으로 정확한 테이블 찾기 (필수)
        기능 2: 서식(폰트, 정렬 등)을 유지하며 데이터 입력 (글자 사라짐 방지)
        """
        data_key = field.get('data_key', '')
        table_config = field.get('table_config', {})
        
        # 설정값 가져오기
        header_row_text = table_config.get('header_row_text', [])
        must_contain = table_config.get('must_contain', [])  # [필수] 테이블 식별 키워드
        columns_config = table_config.get('columns', [])
        
        # 데이터 가져오기
        rows_data = data.get(data_key, [])
        if not rows_data: return
        
        # 빈 행 제거
        rows_data = [r for r in rows_data if isinstance(r, dict) and any(r.values())]
        if not rows_data: return

        self._log(f"   ℹ️ [TABLE_ROWS] '{data_key}' ({len(rows_data)}행) - 표 검색 시작")

        target_table = None
        header_row_idx = None
        norm_headers = [self.normalize_text(h) for h in header_row_text]
        
        # 1. 모든 테이블 순회하며 타겟 찾기
        for t_idx, table in enumerate(all_tables):
            # 컬럼 수 필터 (오차범위 감안)
            if len(table.columns) < len(columns_config) - 1:
                continue

            for r_idx, row in enumerate(table.rows):
                try:
                    cells = row.cells
                    row_texts = [self.normalize_text(c.text) for c in cells]
                    
                    # -----------------------------------------------------------
                    # ✅ [Step 1] 필수 키워드 검사 (Must Contain)
                    # -----------------------------------------------------------
                    if must_contain:
                        is_valid_table = True
                        for req_word in must_contain:
                            norm_req = self.normalize_text(req_word)
                            if not any(norm_req in rt for rt in row_texts):
                                is_valid_table = False
                                break
                        if not is_valid_table:
                            continue  # 필수 단어 없으면 스킵
                    # -----------------------------------------------------------

                    # [Step 2] 헤더 매칭
                    match_count = 0
                    matched_words = []
                    
                    for h in norm_headers:
                        if any(h in rt for rt in row_texts):
                            match_count += 1
                            matched_words.append(h)
                    
                    # 매칭 임계값 (80% 이상)
                    threshold = max(3, int(len(norm_headers) * 0.8))
                    
                    if match_count >= threshold:
                        target_table = table
                        header_row_idx = r_idx
                        self._log(f"      🎯 [타겟 확정] Table #{t_idx}, Row {r_idx}")
                        break
                except Exception as e:
                    continue
            
            if target_table: break
        
        if not target_table:
            self._log(f"      ⚠️ 테이블 못 찾음: {data_key}")
            return

        # 2. 데이터 입력 (서식 유지 로직 적용)
        start_row_offset = table_config.get('start_row_offset', 1)
        data_start_row = header_row_idx + start_row_offset
        max_rows = table_config.get('max_rows', 10)

        for i, row_data in enumerate(rows_data[:max_rows]):
            current_row = data_start_row + i
            
            # 테이블 행 부족 시 자동 추가
            if current_row >= len(target_table.rows):
                self._log(f"      ➕ 행 추가 (Row {current_row})")
                target_table.add_row()
            
            try:
                target_cells = target_table.rows[current_row].cells
                
                for col in columns_config:
                    idx = col.get('col_index')
                    key = col.get('key')
                    val = row_data.get(key, '')
                    
                    if idx < len(target_cells):
                        if val:
                            if hasattr(val, 'strftime'):
                                val = val.strftime('%Y-%m-%d')
                            
                            val_str = str(val)
                            cell = target_cells[idx]
                            
                            # -------------------------------------------------------
                            # ✅ [Step 3] 서식(폰트/크기) 유지하며 값 입력
                            # cell.text = ... 대신 기존 문단/Run을 활용
                            # -------------------------------------------------------
                            if cell.paragraphs:
                                p = cell.paragraphs[0]
                                if p.runs:
                                    # 1) 기존 글자 스타일(Run)이 있으면 텍스트만 교체 (가장 안전)
                                    #    기존 텍스트를 모두 지우고 첫 번째 Run에 새 값 설정
                                    for run in p.runs[1:]: 
                                        run.text = "" # 뒤쪽 찌꺼기 텍스트 제거
                                    p.runs[0].text = val_str
                                else:
                                    # 2) Run은 없지만 문단은 있는 경우 (문단 설정 유지)
                                    p.text = val_str
                            else:
                                # 3) 문단도 없는 깡통 셀인 경우 (어쩔 수 없이 그냥 입력)
                                cell.text = val_str
                            # -------------------------------------------------------

                    else:
                        self._log(f"        ⚠️ Col {idx} 범위 초과")

                self._log(f"        ✅ Row {i+1} 입력 완료")
            except Exception as e:
                self._log(f"        ❌ 입력 오류 (Row {i+1}): {e}")
        
    def _execute_table_cell(self, table, r_idx: int, c_idx: int, 
                            value: Any, field: Dict, current_cell, row_cells):
        """
        TABLE_CELL 전략: 지정된 열 인덱스의 셀에 값 입력
        
        field 설정 예시:
        {
            "data_key": "inviter_name",
            "anchor_text": "성명",
            "strategy": "TABLE_CELL",
            "column_index": 1  # 0-indexed, 초청인=1, 피초청인=2
        }
        """
        column_index = field.get('column_index')
        
        if column_index is None:
            self._log(f"      ⚠️ [TABLE_CELL] column_index 설정 없음")
            return
        
        # 지정된 열 인덱스의 셀에 값 입력
        if column_index < len(row_cells):
            target_cell = row_cells[column_index]
            val_str = str(value)
            
            # 서식 유지하며 입력
            if target_cell.paragraphs:
                p = target_cell.paragraphs[0]
                if p.runs:
                    # 기존 스타일 유지
                    for run in p.runs[1:]:
                        run.text = ""
                    p.runs[0].text = val_str
                else:
                    p.text = val_str
            else:
                target_cell.text = val_str
            
            self._log(f"      ✅ [TABLE_CELL] Row{r_idx}-Col{column_index}에 '{val_str[:20]}...' 입력")
        else:
            self._log(f"      ⚠️ [TABLE_CELL] column_index {column_index} 범위 초과 (max: {len(row_cells)-1})")


    # =============================================================================
    # [전체 수정된 _execute_strategy 메서드]
    # 기존 메서드를 아래로 교체하세요
    # =============================================================================

    def _execute_strategy(self, strategy: str, table, r_idx: int, c_idx: int, 
                        value: Any, field: Dict, current_cell, row_cells):
        """전략별 실행"""
        try:
            # HIERARCHICAL_CHECKBOX
            if strategy == "HIERARCHICAL_CHECKBOX":
                self._execute_hierarchical_checkbox(
                    table, r_idx, c_idx, value, field, current_cell, row_cells
                )
                return
            
            # CHECKBOX_WITH_VALUE
            elif strategy == "CHECKBOX_WITH_VALUE":
                self._execute_checkbox_with_value(
                    table, r_idx, c_idx, value, field, current_cell, row_cells, {}
                )
                return
            
            # ★★★ TABLE_CELL - 새로 추가 ★★★
            elif strategy == "TABLE_CELL":
                self._execute_table_cell(table, r_idx, c_idx, value, field, current_cell, row_cells)
                return
            
            # CHECKBOX
            elif strategy == "CHECKBOX":
                value_map_name = field.get('value_map', 'GENDER')
                value_map = VALUE_MAPS.get(value_map_name, GENDER_CHECKBOX_MAP)
                target_candidates = value_map.get(str(value), [str(value)])
                
                scan_targets = [("현재칸", current_cell)]
                
                for i in range(c_idx + 1, len(row_cells)):
                    scan_targets.append((f"오른쪽+{i-c_idx}", row_cells[i]))
                
                if r_idx + 1 < len(table.rows):
                    try:
                        below_row_cells = table.rows[r_idx + 1].cells
                        for k in range(max(0, c_idx - 1), min(len(below_row_cells), c_idx + 3)):
                            scan_targets.append((f"아래쪽(C{k})", below_row_cells[k]))
                    except:
                        pass
                
                checked_success = False
                
                for pos_name, cell in scan_targets:
                    original_text = cell.text
                    if not original_text.strip():
                        continue
                    
                    for target in target_candidates:
                        pattern = fr"(\[\s*\]|□|☐|\(\s*\))(\s*)({re.escape(target)})"
                        match = re.search(pattern, original_text)
                        
                        if match:
                            new_text = re.sub(pattern, fr"[V]\2\3", original_text, count=1)
                            cell.text = new_text
                            self._log(f"      ✅ [체크] {pos_name}에서 '{target}' 선택")
                            checked_success = True
                            break
                    
                    if checked_success:
                        break
                
                if not checked_success:
                    self._log(f"      ⚠️ [체크 실패] 체크박스 못 찾음: {target_candidates}")
            
            # SPLIT_CELLS
            elif strategy == "SPLIT_CELLS":
                options = field.get('options', {})
                skip_chars = options.get('skip_chars', ['-', '.', ' '])
                val_str = str(value)
                
                for char in skip_chars:
                    val_str = val_str.replace(char, "")
                
                candidates = []
                for i in range(c_idx + 1, len(row_cells)):
                    cell = row_cells[i]
                    if (cell._tc != current_cell._tc) and ("-" not in cell.text):
                        candidates.append(cell)
                
                self._log(f"      ℹ️ [분할] '{val_str}' -> {len(candidates)}칸")
                
                for i, char in enumerate(val_str):
                    if i < len(candidates):
                        candidates[i].text = char
                        for p in candidates[i].paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # NEXT_CELL
            elif strategy == "NEXT_CELL":
                for i in range(c_idx + 1, len(row_cells)):
                    cell = row_cells[i]
                    if cell._tc != current_cell._tc:
                        cell.text = str(value)
                        self._log(f"      ✅ [입력] 오른쪽 셀에 '{str(value)[:20]}...'")
                        return
                
                self._log(f"      ⚠️ [NEXT_CELL 실패] 오른쪽 셀 없음")
            
            # BELOW_CELL
            elif strategy == "BELOW_CELL":
                if r_idx + 1 < len(table.rows):
                    below_cells = table.rows[r_idx + 1].cells
                    if c_idx < len(below_cells):
                        below_cells[c_idx].text = str(value)
                        self._log(f"      ✅ [입력] 아래 셀에 '{str(value)[:20]}...'")
                    else:
                        self._log(f"      ⚠️ [BELOW_CELL 실패] 아래 셀 인덱스 초과")
                else:
                    self._log(f"      ⚠️ [BELOW_CELL 실패] 아래 행 없음")
            
            # APPEND_TO_SAME_CELL
            elif strategy == "APPEND_TO_SAME_CELL":
                val_str = str(value)
                
                if val_str and (val_str not in current_cell.text):
                    if current_cell.text.strip():
                        current_cell.add_paragraph(val_str)
                    else:
                        current_cell.text = val_str
                        
                    self._log(f"      ✅ [추가] 같은 셀(줄바꿈)에 '{val_str[:20]}...'")
        
        except Exception as e:
            self._log(f"      ❌ 처리 중 에러: {e}")
# =============================================================================
# DocumentService 클래스
# =============================================================================

class DocumentService:
    """문서 서비스 클래스"""
    
    def __init__(self, templates_dir: str = "templates"):
        self.templates_dir = templates_dir
        
        if DOCX_AVAILABLE:
            print("✅ python-docx 사용 가능")
        else:
            print("❌ python-docx 미설치")
    
    def get_template_path(self, doc_name: str) -> Optional[str]:
        """문서명으로 템플릿 파일 경로 가져오기"""
        from config.settings import DOCUMENT_TEMPLATES
        
        template_file = DOCUMENT_TEMPLATES.get(doc_name)
        if not template_file:
            print(f"⚠️ 템플릿 매핑 없음: {doc_name}")
            return None
        
        template_path = os.path.join(self.templates_dir, template_file)
        
        if not os.path.exists(template_path):
            print(f"⚠️ 템플릿 파일 없음: {template_path}")
            return None
        
        return template_path
    
    def generate_document(self, doc_name: str, user_data: Dict, 
                         form_data: Dict, narrative_data: Dict) -> bytes:
        """단일 문서 생성"""
        from templates.mapping_guide import get_document_mapping
        
        template_path = self.get_template_path(doc_name)
        
        if not DOCX_AVAILABLE or not template_path:
            return self._create_fallback_document(doc_name, user_data, form_data, narrative_data)
        
        mapping_config = get_document_mapping(doc_name)
        
        if not mapping_config:
            print(f"ℹ️ 매핑 설정 없음: {doc_name}")
            try:
                with open(template_path, 'rb') as f:
                    return f.read()
            except:
                return self._create_fallback_document(doc_name, user_data, form_data, narrative_data)
        
        try:
            processor = DocumentProcessor(user_data, form_data, narrative_data)
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                temp_output = tmp.name
            
            if processor.process_file(mapping_config, template_path, temp_output):
                with open(temp_output, 'rb') as f:
                    doc_bytes = f.read()
                
                try:
                    os.remove(temp_output)
                except:
                    pass
                
                return doc_bytes
            else:
                return self._create_fallback_document(doc_name, user_data, form_data, narrative_data)
                
        except Exception as e:
            print(f"❌ 문서 생성 오류: {str(e)}")
            import traceback
            traceback.print_exc()
            return self._create_fallback_document(doc_name, user_data, form_data, narrative_data)
    
    def _create_fallback_document(self, doc_name: str, user_data: Dict, 
                                  form_data: Dict, narrative_data: Dict) -> bytes:
        """폴백 텍스트 문서 생성"""
        lines = []
        lines.append("=" * 60)
        lines.append(f"  {doc_name}")
        lines.append("=" * 60)
        lines.append("")
        
        lines.append("[신청인 정보 - Layer 1]")
        if user_data:
            for key, value in user_data.items():
                if value:
                    lines.append(f"  {key}: {value}")
        else:
            lines.append("  (정보 없음)")
        lines.append("")
        
        lines.append("[시나리오별 정보 - Layer 2]")
        if form_data:
            for key, value in form_data.items():
                if value:
                    # 배열인 경우 (TABLE_ROWS 데이터)
                    if isinstance(value, list):
                        lines.append(f"  {key}: [{len(value)}개 행]")
                        for idx, row in enumerate(value):
                            lines.append(f"    Row {idx+1}: {row}")
                    else:
                        lines.append(f"  {key}: {value}")
        else:
            lines.append("  (정보 없음)")
        lines.append("")
        
        lines.append("[서술형 내용 - Layer 3]")
        if narrative_data:
            for key, value in narrative_data.items():
                if value:
                    lines.append(f"  {key}: {value}")
        else:
            lines.append("  (내용 없음)")
        lines.append("")
        
        lines.append("=" * 60)
        lines.append(f"  생성일시: {datetime.now().strftime('%Y년 %m월 %d일 %H:%M')}")
        lines.append("  K-Stay - Korea Stay Assistant")
        lines.append("=" * 60)
        
        return "\n".join(lines).encode('utf-8')
    
    def generate_full_package(self, scenario_id: str, user_data: Dict,
                             form_data: Dict, narrative_data: Dict) -> bytes:
        """시나리오별 전체 문서 패키지 생성 (ZIP)"""
        from config.settings import SCENARIOS
        from templates.mapping_guide import get_scenario_documents
        
        scenario = SCENARIOS.get(scenario_id)
        if not scenario:
            st.error("유효하지 않은 시나리오입니다.")
            return b""
        
        required_docs = get_scenario_documents(scenario_id)
        if not required_docs:
            required_docs = scenario.required_docs
        
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for doc_name in required_docs:
                try:
                    doc_bytes = self.generate_document(
                        doc_name, user_data, form_data, narrative_data
                    )
                    
                    safe_name = doc_name.replace(' ', '_').replace('/', '_')
                    
                    if len(doc_bytes) >= 4 and doc_bytes[:4] == b'PK\x03\x04':
                        filename = f"{safe_name}.docx"
                    else:
                        filename = f"{safe_name}.txt"
                    
                    zip_file.writestr(filename, doc_bytes)
                    print(f"📄 추가됨: {filename}")
                    
                except Exception as e:
                    error_content = f"문서 생성 오류: {str(e)}"
                    zip_file.writestr(f"ERROR_{doc_name}.txt", error_content.encode('utf-8'))
                    print(f"❌ 오류: {doc_name} - {str(e)}")
            
            readme_content = self._create_readme(scenario, required_docs, datetime.now())
            zip_file.writestr("README.txt", readme_content.encode('utf-8'))
        
        zip_buffer.seek(0)
        return zip_buffer.getvalue()
    
    def _create_readme(self, scenario, docs: List[str], generated_at: datetime) -> str:
        """README 파일 생성"""
        lines = [
            "=" * 60,
            "K-Stay Document Package",
            "=" * 60,
            "",
            f"시나리오: {scenario.name} ({scenario.visa_type})",
            f"생성일시: {generated_at.strftime('%Y년 %m월 %d일 %H:%M:%S')}",
            "",
            "포함된 문서:",
            "-" * 40,
        ]
        
        for i, doc in enumerate(docs, 1):
            lines.append(f"  {i}. {doc}")
        
        lines.extend([
            "",
            "=" * 60,
            "Powered by K-Stay",
            "=" * 60,
        ])
        
        return "\n".join(lines)


# =============================================================================
# DocumentPreviewService 클래스
# =============================================================================

class DocumentPreviewService:
    """문서 미리보기 서비스"""
    
    @staticmethod
    def preview_document(doc_bytes: bytes, doc_name: str):
        """문서 미리보기 렌더링"""
        try:
            try:
                content = doc_bytes.decode('utf-8')
                st.markdown(f"""
                    <div style="
                        background: #f8fafc;
                        border: 1px solid #e2e8f0;
                        border-radius: 12px;
                        padding: 1.5rem;
                        font-family: monospace;
                        font-size: 0.9rem;
                        white-space: pre-wrap;
                        max-height: 500px;
                        overflow-y: auto;
                    ">
{content}
                    </div>
                """, unsafe_allow_html=True)
            except UnicodeDecodeError:
                st.info(f"📄 {doc_name} - Word 문서 파일입니다. 다운로드하여 확인하세요.")
                
        except Exception as e:
            st.error(f"미리보기 오류: {str(e)}")
    
    @staticmethod
    def render_download_section(zip_bytes: bytes, scenario_name: str):
        """다운로드 섹션 렌더링"""
        
        st.markdown("""
            <div style="
                background: linear-gradient(135deg, #eff6ff, #dbeafe);
                border: 2px solid #3b82f6;
                border-radius: 20px;
                padding: 3rem;
                text-align: center;
                margin: 2rem 0;
            ">
                <h2 style="color: #1d4ed8; margin-bottom: 1rem;">
                    📦 문서 패키지 준비 완료!
                </h2>
                <p style="color: #64748b; margin-bottom: 2rem;">
                    모든 문서가 성공적으로 생성되었습니다.
                </p>
            </div>
        """, unsafe_allow_html=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"KStay_{scenario_name}_{timestamp}.zip"
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.download_button(
                label="📥 ZIP 패키지 다운로드",
                data=zip_bytes,
                file_name=filename,
                mime="application/zip",
                use_container_width=True
            )